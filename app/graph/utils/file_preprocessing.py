from pathlib import Path
from dataclasses import dataclass
import re
from docx import Document


@dataclass
class CVData:
    """CV raw data"""

    text: str
    original_filename: str
    word_count: int


class CVPreprocessingError(Exception):
    pass


class CVPreprocessor:
    """
    DOCX→text

    """

    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    ALLOWED_EXTENSIONS = {".docx"}

    def process(self, file_path: str) -> CVData:
        """
        Waliduj + parsuj DOCX do tekstu

        Args:
            file_path: Ścieżka do pliku DOCX

        Returns:
            CVData: Obiekt z czystym tekstem

        Raises:
            CVPreprocessingError: W przypadku błędów
        """
        path = Path(file_path)

        # 1. Walidacja
        self._validate_file(path)

        # 2. Parsowanie DOCX → text
        try:
            document = Document(str(path))
            text = self._extract_text(document)
        except Exception as e:
            raise CVPreprocessingError(f"Błąd parsowania DOCX: {str(e)}")

        # 3. Podstawowe czyszczenie
        text = self._clean_text(text)

        # 4. Sprawdź czy jest jakaś treść
        if not text or len(text.strip()) < 50:
            raise CVPreprocessingError("Plik nie zawiera wystarczającej ilości tekstu")

        return CVData(
            text=text, original_filename=path.name, word_count=len(text.split())
        )

    def _validate_file(self, path: Path) -> None:
        """Walidacja techniczna pliku"""

        if not path.exists():
            raise CVPreprocessingError(f"Plik nie istnieje: {path}")

        if not path.is_file():
            raise CVPreprocessingError(f"Ścieżka nie wskazuje na plik: {path}")

        if path.suffix.lower() not in self.ALLOWED_EXTENSIONS:
            raise CVPreprocessingError(
                f"Nieprawidłowe rozszerzenie: {path.suffix}. "
                f"Dozwolone: {', '.join(self.ALLOWED_EXTENSIONS)}"
            )

        file_size = path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            raise CVPreprocessingError(
                f"Plik za duży: {file_size / 1024 / 1024:.2f}MB. "
                f"Max: {self.MAX_FILE_SIZE / 1024 / 1024:.0f}MB"
            )

        if file_size == 0:
            raise CVPreprocessingError("Plik jest pusty")

    def _extract_text(self, document: Document) -> str:
        """
        Ekstrakcja tekstu z DOCX (paragrafy + tabele)
        Zachowuje naturalną kolejność elementów
        """
        text_parts = []

        # Paragrafy
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Tabele
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        return "\n".join(text_parts)

    def _clean_text(self, text: str) -> str:
        """
        Minimalne czyszczenie - tylko oczywiste problemy
        LLM poradzi sobie z resztą
        """
        # Usuń nadmiarowe spacje
        text = re.sub(r" +", " ", text)

        # Usuń więcej niż 2 puste linie pod rząd
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Trim
        text = text.strip()

        return text
